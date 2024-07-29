#core.word parser
from docx import Document
import re

class WordParser:
    def __init__(self, word_path):
        self.document = Document(word_path)

    def get_placeholders(self):
        placeholders = set()
        for paragraph in self.document.paragraphs:
            placeholders.update(re.findall(r'\{\{(\w+)\}\}', paragraph.text))
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    placeholders.update(re.findall(r'\{\{(\w+)\}\}', cell.text))
        print(f"Placeholders found: {placeholders}")  # Add this line
        return list(placeholders)
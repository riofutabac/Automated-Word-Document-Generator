# core/document_generator.py
# core/document_generator.py
from docx import Document

class DocumentGenerator:
    def __init__(self, template_path):
        self.template_path = template_path

    def generate_document(self, data):
        doc = Document(self.template_path)
        for p in doc.paragraphs:
            self.replace_placeholders(p, data)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self.replace_placeholders(p, data)
        return doc

    def replace_placeholders(self, paragraph, data):
        for key, value in data.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, str(value))
                # Esto reemplaza el texto completo del párrafo y asegura que el placeholder es sustituido

    def save_document(self, doc, output_path):
        doc.save(output_path)
